"""Resume-focused council flow."""

from typing import List, Dict, Any, Tuple, Optional
import re
import base64
from io import BytesIO
from docx import Document
from docx.shared import Pt

from .openrouter import query_models_parallel, query_model
from .config import (
    RESUME_DRAFT_MODELS,
    RESUME_RANKING_MODELS,
    RESUME_JUDGE_MODEL,
    RESUME_POLISH_MODEL,
    RESUME_POLISH_THRESHOLD,
    RESUME_DRAFT_MAX_TOKENS,
    RESUME_JUDGE_MAX_TOKENS,
    RESUME_POLISH_MAX_TOKENS,
    RESUME_USE_PEER_RANKING,
    RESUME_PROFILE_PACK_MAX_CHARS,
    RESUME_SEND_FULL_PROFILE,
    RESUME_STYLE_GUIDE,
    RESUME_STYLE_GUIDE_PATH,
)
from .council import parse_ranking_from_text, calculate_aggregate_rankings
from .packs import build_profile_pack, build_jd_pack, basic_resume_heuristics, extract_profile_section


_REQUIRED_SECTIONS = [
    "Summary",
    "Education",
    "Technical Skills",
    "Professional Experience",
    "Projects",
    "Certifications",
]


def _load_style_guide_text() -> str:
    if isinstance(RESUME_STYLE_GUIDE, str) and RESUME_STYLE_GUIDE.strip():
        return RESUME_STYLE_GUIDE.strip()
    path = (RESUME_STYLE_GUIDE_PATH or "").strip()
    if not path:
        return ""
    try:
        with open(path, "r", encoding="utf-8") as f:
            return (f.read() or "").strip()
    except Exception:
        return ""


def _default_resume_style_guide() -> str:
    return """You are an expert UK-based CV writer for AI/Data/Legal Tech/regulatory roles.

Hard formatting rules:
- Resume tone: third-person / action-verb based. NEVER first-person pronouns (I/me/my).
- Summary: 2–3 sentences MAX. No narrative storytelling.
- Bullets: short, punchy, 1–2 lines; avoid multi-clause explanations (no “so that/because/in order to”).
- Experience: 3–4 bullets per role.
- Projects: include ONLY 2–3 most relevant; each is ONE-LINER: **Project Name** | Tech Stack | 1-sentence outcome.
- British English; no tables/graphics.

Truth/metrics guardrails:
- Do NOT invent facts.
- Only include exact metrics (%, counts, time saved) if verbatim in MASTER PROFILE or JD.
- If unsure, convert to qualitative wording.
"""


def _resume_style_block() -> str:
    user_guide = _load_style_guide_text()
    if user_guide:
        return _default_resume_style_guide() + "\n\n---\nCUSTOM STRATEGY DOC (authoritative):\n" + user_guide
    return _default_resume_style_guide()


def _normalize_heading(line: str) -> str:
    s = (line or "").strip()
    if s.startswith("#"):
        s = s.lstrip("#").strip()
    s = s.rstrip(":").rstrip("-").rstrip("–").rstrip("—").strip()
    return s.lower()


def _split_resume_sections(markdown_text: str) -> Dict[str, List[str]]:
    lines = (markdown_text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    sections: Dict[str, List[str]] = {}

    current: Optional[str] = None
    for line in lines:
        norm = _normalize_heading(line)
        matched = next((h for h in _REQUIRED_SECTIONS if norm == h.lower()), None)
        if matched:
            current = matched
            sections.setdefault(current, [])
            continue
        if current is not None:
            sections[current].append(line.rstrip())

    return sections


def _is_effectively_empty(lines: List[str]) -> bool:
    content = "\n".join([l for l in (lines or []) if (l or "").strip()]).strip()
    if not content:
        return True
    if content.strip().lower() == "n/a":
        return True
    return False


def _coerce_bullets(lines: List[str], max_lines: int = 8) -> List[str]:
    out: List[str] = []
    for line in (lines or []):
        s = (line or "").strip()
        if not s:
            continue
        # normalize common bullet markers
        if s.startswith("•"):
            s = "- " + s.lstrip("•").strip()
        if s.startswith("*") and not s.startswith("**"):
            s = "- " + s.lstrip("*").strip()
        # Clamp very long bullets to reduce multi-line spillover.
        if len(s) > 180:
            s = s[:177].rstrip() + "..."
        out.append(s)
        if len(out) >= max_lines:
            break
    return out


def _ensure_required_outline(markdown_text: str, profile_pack: str) -> str:
    sections = _split_resume_sections(markdown_text)

    # Fill missing / empty sections.
    for heading in _REQUIRED_SECTIONS:
        if heading not in sections:
            sections[heading] = ["N/A"]

    # If Projects/Certifications are empty, try to pull from truth source.
    for heading in ("Projects", "Certifications"):
        if not _is_effectively_empty(sections.get(heading, [])):
            continue

        extracted = (extract_profile_section(profile_pack, heading) or "").strip()
        if extracted:
            extracted_lines = extracted.splitlines()
            # drop the heading line itself
            extracted_lines = extracted_lines[1:] if len(extracted_lines) > 1 else []
            bullets = _coerce_bullets(extracted_lines, max_lines=3)
            if bullets:
                sections[heading] = bullets
                continue

        sections[heading] = ["N/A"]

    # Rebuild in strict order.
    rebuilt: List[str] = []
    for heading in _REQUIRED_SECTIONS:
        rebuilt.append(heading)
        content_lines = [l for l in (sections.get(heading) or [])]
        # Ensure at least one non-empty line
        if _is_effectively_empty(content_lines):
            content_lines = ["N/A"]
        # Trim leading/trailing blank lines
        while content_lines and not (content_lines[0] or "").strip():
            content_lines = content_lines[1:]
        while content_lines and not (content_lines[-1] or "").strip():
            content_lines = content_lines[:-1]
        # Remove obvious truncation artifacts / dangling markdown markers.
        content_lines = [l for l in content_lines if (l or "").strip() not in {"**", "*"}]
        rebuilt.extend(content_lines)
        rebuilt.append("")

    return "\n".join(rebuilt).strip() + "\n"


def _missing_required_sections(markdown_text: str) -> bool:
    lower = (markdown_text or "").lower()
    return any(h.lower() not in lower for h in _REQUIRED_SECTIONS)


def _resume_prompt(profile_pack: str, jd_pack: Dict[str, object], company_details: str) -> str:
    keywords = ", ".join(jd_pack.get("keywords", []))
    jd_compact = jd_pack.get("jd_compact", "")
    return f"""{_resume_style_block()}

You are a resume writer. Produce a concise resume in strict markdown with these exact sections and order.

    Hard rule: NEVER omit any of these headings. If a section has no items, keep the heading and put "N/A" on the next line.

Summary
Education
Technical Skills
Professional Experience
Projects
Certifications

Formatting rules:
- Use short, impact bullets; no fluff; no invented facts.
- Bold project names.
- Section headings exactly as above.
- Title corrections (hard rules):
    - Any Venrag experience title must be "AI Engineer" (even if source says Research Developer or Data Scientist).
    - Any Elite experience title must be "Data Administrator" (not Admin).
    - Any CVA experience title must be "Database Administrator".
- Keep content truthful to MASTER PROFILE.

MASTER PROFILE (truth source):
{profile_pack}

JOB DESCRIPTION (target role):
{jd_compact}

ATS KEYWORDS (prioritize when truthful):
{keywords}

COMPANY DETAILS (tone/culture hints):
{company_details}

Return only the resume markdown, nothing else."""


def _peer_ranking_prompt(resume_blocks: List[Tuple[str, str]], jd_pack: Dict[str, object], profile_pack: str) -> str:
    merged = "\n\n".join([f"{label}:\n{resume}" for label, resume in resume_blocks])
    return f"""{_resume_style_block()}

You are ranking tailored resumes for a job.

TRUTH SOURCE (Profile Facts Pack):
{profile_pack}

JOB REQUIREMENTS PACK:
{jd_pack}

Resumes (anonymized):
{merged}

Evaluate each resume on:
- Keyword coverage vs JD
- Relevance/alignment to role and company
- Truthfulness to the profile pack (no invented claims)
- Formatting/clarity

Provide brief feedback per resume. Then give FINAL RANKING as numbered list using only the labels (e.g., "1. Response A")."""


def _normalize_parsed_ranking(parsed: List[str], label_to_model: Dict[str, str]) -> List[str]:
    """De-duplicate parsed labels and keep only known labels, preserving order."""
    out: List[str] = []
    seen = set()
    for label in (parsed or []):
        if label not in label_to_model:
            continue
        if label in seen:
            continue
        seen.add(label)
        out.append(label)
    return out


def _best_model_from_metadata(metadata: Dict[str, Any]) -> Optional[str]:
    agg = (metadata or {}).get("aggregate_rankings") or []
    if isinstance(agg, list) and agg:
        top = agg[0] or {}
        model = top.get("model")
        if isinstance(model, str) and model.strip():
            return model
    return None


def _judge_prompt(
    resume_blocks: List[Tuple[str, str]],
    profile_pack: str,
    jd_pack: Dict[str, object],
    heuristics: Dict[str, Dict[str, object]],
) -> str:
    merged = "\n\n".join([f"{label}:\n{resume}" for label, resume in resume_blocks])
    return f"""{_resume_style_block()}

You are judging anonymized resume drafts for a specific job.

TRUTH SOURCE (Profile Facts Pack):
{profile_pack}

JOB REQUIREMENTS PACK:
{jd_pack}

Cheap heuristics (computed by code):
{heuristics}

Resumes:
{merged}

Task:
1) Score each resume 0-100 on: keyword_coverage, role_relevance, truthfulness, formatting.
2) Output STRICT JSON with this schema:
{{
  "scores": [{{"label":"Response A","keyword_coverage":0,"role_relevance":0,"truthfulness":0,"formatting":0,"overall":0,"notes":"..."}}],
  "winner": "Response A",
  "final_ranking": ["Response A","Response B"],
  "unsupported_claims": ["..."]
}}
3) After the JSON, include a FINAL RANKING section formatted exactly:
FINAL RANKING:
1. Response A
2. Response B
"""


def _chairman_prompt(user_query: str, stage1_results: List[Dict[str, Any]], stage2_results: List[Dict[str, Any]]) -> str:
    stage1_text = "\n\n".join([
        f"Model: {res['model']}\nResume:\n{res['response']}" for res in stage1_results
    ])
    stage2_text = "\n\n".join([
        f"Model: {res['model']}\nRanking:\n{res['ranking']}" for res in stage2_results
    ])
    return f"""You are the chairman finalizing the best truthful resume.

Input question/context: {user_query}

Stage 1 resumes:
{stage1_text}

Stage 2 rankings and feedback:
{stage2_text}

Produce a final, polished resume in markdown that:
- Stays 100% faithful to the master profile facts
- Maximizes alignment with the job description
- Uses strong, concise bullets and keywords
Return the resume markdown only."""


async def stage1_generate_resumes(profile_pack: str, jd_pack: Dict[str, object], company_details: str) -> List[Dict[str, Any]]:
    messages = [{"role": "user", "content": _resume_prompt(profile_pack, jd_pack, company_details)}]
    responses = await query_models_parallel(
        RESUME_DRAFT_MODELS,
        messages,
        timeout=90.0,
        max_tokens=RESUME_DRAFT_MAX_TOKENS,
        temperature=0.5,
    )
    results = []
    for model, response in responses.items():
        if response is None:
            continue
        content = (response.get("content", "") or "").strip()
        if not content:
            continue

        # Normalize to the required outline so downstream ranking + DOCX is stable.
        content = _ensure_required_outline(content, profile_pack)
        results.append({
            "model": model,
            # Keep API consistent with existing UI components (Stage1/Stage3 expect 'response')
            "response": content
        })
    return results


async def stage2_judge_resumes(profile_pack: str, jd_pack: Dict[str, object], stage1_results: List[Dict[str, Any]]):
    labels = [chr(65 + i) for i in range(len(stage1_results))]
    label_to_model = {f"Response {label}": res["model"] for label, res in zip(labels, stage1_results)}
    resume_blocks = [(f"Response {label}", res["response"]) for label, res in zip(labels, stage1_results)]

    heuristics = {
        label: basic_resume_heuristics(resume, jd_pack.get("keywords", []))
        for (label, resume) in resume_blocks
    }

    messages = [{"role": "user", "content": _judge_prompt(resume_blocks, profile_pack, jd_pack, heuristics)}]
    response = await query_model(
        RESUME_JUDGE_MODEL,
        messages,
        timeout=60.0,
        max_tokens=RESUME_JUDGE_MAX_TOKENS,
        temperature=0.2,
    )

    full_text = (response or {}).get("content", "")
    parsed = _normalize_parsed_ranking(parse_ranking_from_text(full_text), label_to_model)
    rankings = [
        {
            "model": RESUME_JUDGE_MODEL,
            "ranking": full_text,
            "parsed_ranking": parsed,
        }
    ]

    # For UI reuse, create an aggregate list from the judge ranking only.
    aggregate_rankings = []
    for idx, label in enumerate(parsed, start=1):
        if label in label_to_model:
            aggregate_rankings.append(
                {
                    "model": label_to_model[label],
                    "average_rank": float(idx),
                    "rankings_count": 1,
                }
            )

    metadata = {"label_to_model": label_to_model, "aggregate_rankings": aggregate_rankings, "heuristics": heuristics}
    return rankings, metadata


async def stage2_peer_rank_resumes(profile_pack: str, jd_pack: Dict[str, object], stage1_results: List[Dict[str, Any]]):
    labels = [chr(65 + i) for i in range(len(stage1_results))]
    label_to_model = {f"Response {label}": res["model"] for label, res in zip(labels, stage1_results)}
    resume_blocks = [(f"Response {label}", res["response"]) for label, res in zip(labels, stage1_results)]

    # Require strict FINAL RANKING format so parsing works.
    strict_prompt = _peer_ranking_prompt(resume_blocks, jd_pack, profile_pack) + "\n\nIMPORTANT: End with a section exactly:\nFINAL RANKING:\n1. Response A\n2. Response B\n..."
    messages = [{"role": "user", "content": strict_prompt}]
    responses = await query_models_parallel(
        RESUME_RANKING_MODELS,
        messages,
        timeout=90.0,
        max_tokens=RESUME_JUDGE_MAX_TOKENS,
        temperature=0.2,
    )

    rankings = []
    for model, response in responses.items():
        if response is None:
            continue
        full_text = (response.get("content", "") or "").strip()
        if not full_text:
            continue
        parsed = _normalize_parsed_ranking(parse_ranking_from_text(full_text), label_to_model)
        rankings.append({
            "model": model,
            "ranking": full_text,
            "parsed_ranking": parsed,
        })

    aggregate = calculate_aggregate_rankings(rankings, label_to_model)
    metadata = {"label_to_model": label_to_model, "aggregate_rankings": aggregate}
    return rankings, metadata


def _polish_prompt(profile_pack: str, jd_pack: Dict[str, object], company_details: str, best_resume: str) -> str:
    return f"""{_resume_style_block()}

Polish this resume for the job while staying 100% truthful to the profile pack.

PROFILE FACTS PACK:
{profile_pack}

JOB REQUIREMENTS PACK:
{jd_pack}

COMPANY DETAILS:
{company_details}

DRAFT RESUME (markdown):
{best_resume}

Rules:
- Keep the same required sections: Summary, Education, Technical Skills, Professional Experience, Projects, Certifications
- Bullet points only under experience/projects.
- Bold project names.
- Do NOT invent facts.

Return only the improved resume markdown."""


def _pick_best_label(stage2_results: List[Dict[str, Any]]) -> Optional[str]:
    if not stage2_results:
        return None
    parsed = stage2_results[0].get("parsed_ranking") or []
    if parsed:
        return parsed[0]
    return None


async def stage3_finalize(profile_pack: str, jd_pack: Dict[str, object], company_details: str, stage1_results: List[Dict[str, Any]], stage2_results: List[Dict[str, Any]], metadata: Dict[str, Any]):
    label_to_model = (metadata or {}).get("label_to_model") or {}
    best_model = _best_model_from_metadata(metadata)
    best_label = None
    if best_model:
        best_resume = next((r["response"] for r in stage1_results if r["model"] == best_model), "")
    else:
        best_label = _pick_best_label(stage2_results)
        if best_label and best_label in label_to_model:
            best_model = label_to_model[best_label]
            best_resume = next((r["response"] for r in stage1_results if r["model"] == best_model), "")
        else:
            best_resume = stage1_results[0].get("response", "")
            best_model = stage1_results[0].get("model", "unknown")

    if not (best_resume or "").strip():
        # Last resort: pick the first non-empty draft
        first_non_empty = next((r for r in stage1_results if (r.get("response") or "").strip()), None)
        if first_non_empty:
            best_resume = first_non_empty.get("response", "")
            best_model = first_non_empty.get("model", best_model)

    # Always normalize outline and force polish when drafts are incomplete/truncated.
    raw_missing = _missing_required_sections(best_resume)
    raw_looks_truncated = bool(re.search(r"\*\*$", (best_resume or "").strip()))
    best_resume = _ensure_required_outline(best_resume, profile_pack)

    # Gate premium polish based on heuristics score (cheap proxy), but override when incomplete.
    heur = basic_resume_heuristics(best_resume, jd_pack.get("keywords", []))
    score_proxy = 0.5 * float(heur.get("keyword_hit_rate", 0.0)) + 0.5 * float(heur.get("section_completeness", 0.0))

    if (not raw_missing) and (not raw_looks_truncated) and score_proxy >= RESUME_POLISH_THRESHOLD:
        return {"model": best_model, "response": best_resume, "notes": "Selected best draft (no premium polish)."}

    messages = [{"role": "user", "content": _polish_prompt(profile_pack, jd_pack, company_details, best_resume)}]
    response = await query_model(
        RESUME_POLISH_MODEL,
        messages,
        timeout=90.0,
        max_tokens=RESUME_POLISH_MAX_TOKENS,
        temperature=0.4,
    )
    polished = ((response or {}).get("content", "") or "").strip()
    if response is None or not polished:
        return {"model": best_model, "response": best_resume, "notes": "Premium polish failed; returning best draft."}
    polished = _ensure_required_outline(polished, profile_pack)
    return {"model": RESUME_POLISH_MODEL, "response": polished, "notes": "Premium polish applied."}


def _markdown_to_docx_bytes(markdown_text: str) -> str:
    doc = Document()

    # Set base font to Times New Roman 10
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)

    heading_style = doc.styles["Heading 2"]
    heading_font = heading_style.font
    heading_font.name = "Times New Roman"
    heading_font.size = Pt(10)
    heading_font.bold = True

    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Heading detection (e.g., "Summary" or "Summary:")
        if stripped.endswith(":") or stripped in {
            "Summary", "Education", "Technical Skills", "Professional Experience", "Projects", "Certifications"
        }:
            doc.add_paragraph(stripped.rstrip(":"), style=heading_style)
            continue

        # Bullet lines
        if stripped.startswith("- "):
            content = stripped[2:]
            para = doc.add_paragraph(style="List Bullet")
            # If bolded project name pattern **Name** rest
            if content.startswith("**") and "**" in content[2:]:
                closing = content.find("**", 2)
                name = content[2:closing]
                rest = content[closing+2:].lstrip()
                run_name = para.add_run(name)
                run_name.bold = True
                if rest:
                    para.add_run(f" - {rest}")
            else:
                para.add_run(content)
            continue

        # Default paragraph
        doc.add_paragraph(stripped)

    buf = BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode("utf-8")


async def run_resume_council(
    master_profile: str,
    job_description: str,
    company_details: str,
    use_peer_ranking: Optional[bool] = None,
) -> Tuple[List, List, Dict, Dict, Dict]:
    profile_pack = build_profile_pack(
        master_profile,
        max_chars=None if RESUME_SEND_FULL_PROFILE else RESUME_PROFILE_PACK_MAX_CHARS,
    )
    jd_pack = build_jd_pack(job_description)

    stage1_results = await stage1_generate_resumes(profile_pack, jd_pack, company_details)
    if not stage1_results:
        return [], [], {"model": "error", "response": "No models responded."}, {}, {"base64": "", "filename": "resume.docx"}

    effective_peer_ranking = RESUME_USE_PEER_RANKING if use_peer_ranking is None else bool(use_peer_ranking)

    if effective_peer_ranking:
        stage2_results, metadata = await stage2_peer_rank_resumes(profile_pack, jd_pack, stage1_results)
    else:
        stage2_results, metadata = await stage2_judge_resumes(profile_pack, jd_pack, stage1_results)
    stage3_result = await stage3_finalize(profile_pack, jd_pack, company_details, stage1_results, stage2_results, metadata)

    metadata["profile_pack_chars"] = len(profile_pack)
    metadata["profile_pack_full"] = bool(RESUME_SEND_FULL_PROFILE)
    metadata["jd_pack"] = jd_pack
    metadata["peer_ranking_used"] = effective_peer_ranking
    metadata["draft_models_requested"] = list(RESUME_DRAFT_MODELS)
    metadata["draft_models_returned"] = [r.get("model") for r in (stage1_results or [])]
    returned_set = {m for m in metadata["draft_models_returned"] if isinstance(m, str)}
    metadata["draft_models_missing"] = [m for m in RESUME_DRAFT_MODELS if m not in returned_set]

    docx_base64 = _markdown_to_docx_bytes(stage3_result.get("response", ""))
    docx_info = {"base64": docx_base64, "filename": "tailored_resume.docx"}

    return stage1_results, stage2_results, stage3_result, metadata, docx_info
